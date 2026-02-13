"""Word document (.docx) generation using python-docx."""

from __future__ import annotations

import tempfile
from datetime import datetime, timezone
from io import BytesIO
from urllib.request import urlopen

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from app.models.schemas import BrandConfig, ReportRequest, ReportType, TemplateDesign

_HEADER_BG = RGBColor(0x1E, 0x29, 0x3B)
_HEADER_FG = RGBColor(0xFF, 0xFF, 0xFF)
_ALT_ROW_BG = RGBColor(0xF8, 0xF9, 0xFA)  # Fixed light gray — never brand-neutral
_BORDER_COLOR = RGBColor(0xE5, 0xE7, 0xEB)  # Thinner, lighter borders


def _hex_to_rgb(hex_color: str) -> RGBColor:
    """Convert a hex color string like '#1e293b' to an RGBColor."""
    h = hex_color.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _get_theme_config(req: ReportRequest) -> dict:
    """Return theme-aware styling config based on template_design."""
    bc = req.brand_config or BrandConfig()
    design = req.template_design

    if design == TemplateDesign.PLAIN:
        return {
            "header_bg": RGBColor(0xF9, 0xFA, 0xFB),
            "header_fg": RGBColor(0x37, 0x41, 0x51),
            "alt_row_bg": _ALT_ROW_BG,
            "border_color": _BORDER_COLOR,
            "secondary": RGBColor(0x37, 0x41, 0x51),
            "show_logo": False,
            "use_zebra": False,
        }
    elif design == TemplateDesign.EXECUTIVE:
        return {
            "header_bg": RGBColor(0xF8, 0xFA, 0xFC),
            "header_fg": _hex_to_rgb(bc.primary_color),
            "alt_row_bg": RGBColor(0xFA, 0xFB, 0xFC),
            "border_color": _BORDER_COLOR,
            "secondary": _hex_to_rgb(bc.secondary_color),
            "show_logo": True,
            "use_zebra": True,
        }
    else:  # classic — uses fixed light gray for zebra, NOT brand-neutral
        return {
            "header_bg": _hex_to_rgb(bc.primary_color),
            "header_fg": _HEADER_FG,
            "alt_row_bg": _ALT_ROW_BG,
            "border_color": _BORDER_COLOR,
            "secondary": _hex_to_rgb(bc.secondary_color),
            "show_logo": True,
            "use_zebra": True,
        }


def _get_colors(req: ReportRequest) -> dict[str, RGBColor]:
    """Return color dict from brand_config or fall back to module defaults."""
    return _get_theme_config(req)


def _set_cell_shading(cell, color: RGBColor):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), str(color))
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def _style_header_row(row, columns: list[str], colors: dict[str, RGBColor] | None = None):
    bg = (colors or {}).get("header_bg", _HEADER_BG)
    fg = (colors or {}).get("header_fg", _HEADER_FG)
    for i, cell in enumerate(row.cells):
        cell.text = columns[i]
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = fg
        _set_cell_shading(cell, bg)


def _add_data_row(table, values: list[str], row_idx: int, colors: dict | None = None):
    alt_bg = (colors or {}).get("alt_row_bg", _ALT_ROW_BG)
    use_zebra = (colors or {}).get("use_zebra", True)
    row = table.add_row()
    for i, cell in enumerate(row.cells):
        cell.text = str(values[i])
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(9)
    if use_zebra and row_idx % 2 == 0:
        for cell in row.cells:
            _set_cell_shading(cell, alt_bg)
    return row


def _add_summary_section(doc: Document, items: list[tuple[str, str]], colors: dict[str, RGBColor] | None = None):
    heading_color = (colors or {}).get("header_bg", _HEADER_BG)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("Summary")
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = heading_color

    for label, value in items:
        p = doc.add_paragraph()
        run_label = p.add_run(f"{label}: ")
        run_label.font.bold = True
        run_label.font.size = Pt(10)
        run_val = p.add_run(str(value))
        run_val.font.size = Pt(10)


def _build_staff_shifts(doc: Document, req: ReportRequest, colors: dict[str, RGBColor] | None = None):
    cols = ["Date", "Event", "Client", "Venue", "Role", "Clock In", "Clock Out", "Hours", "Rate", "Earnings"]
    table = doc.add_table(rows=1, cols=len(cols))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _style_header_row(table.rows[0], cols, colors)

    for idx, r in enumerate(req.records):
        _add_data_row(
            table,
            [
                r.get("date", ""),
                r.get("eventName", ""),
                r.get("clientName", ""),
                r.get("venueName", ""),
                r.get("role", ""),
                r.get("clockIn", ""),
                r.get("clockOut", ""),
                str(r.get("hoursWorked", 0)),
                f"${r.get('hourlyRate', 0):,.2f}",
                f"${r.get('earnings', 0):,.2f}",
            ],
            idx,
            colors,
        )

    _add_summary_section(
        doc,
        [
            ("Total Shifts", str(req.summary.get("totalShifts", len(req.records)))),
            ("Total Hours", str(req.summary.get("totalHours", 0))),
            ("Total Earnings", f"${req.summary.get('totalEarnings', 0):,.2f}"),
        ],
        colors,
    )


def _build_payroll(doc: Document, req: ReportRequest, colors: dict[str, RGBColor] | None = None):
    cols = ["Staff Name", "Email", "Shifts", "Hours", "Avg Rate", "Total Pay"]
    table = doc.add_table(rows=1, cols=len(cols))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _style_header_row(table.rows[0], cols, colors)

    for idx, r in enumerate(req.records):
        _add_data_row(
            table,
            [
                r.get("name", ""),
                r.get("email", ""),
                str(r.get("shifts", 0)),
                str(r.get("hours", 0)),
                f"${r.get('averageRate', 0):,.2f}",
                f"${r.get('totalPay', 0):,.2f}",
            ],
            idx,
            colors,
        )

    _add_summary_section(
        doc,
        [
            ("Staff Count", str(req.summary.get("staffCount", len(req.records)))),
            ("Total Hours", str(req.summary.get("totalHours", 0))),
            ("Total Payroll", f"${req.summary.get('totalPayroll', 0):,.2f}"),
        ],
        colors,
    )


def _build_attendance(doc: Document, req: ReportRequest, colors: dict[str, RGBColor] | None = None):
    cols = ["Date", "Event", "Staff", "Role", "Sched. Start", "Sched. End", "Clock In", "Clock Out", "Hours", "Status"]
    table = doc.add_table(rows=1, cols=len(cols))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _style_header_row(table.rows[0], cols, colors)

    for idx, r in enumerate(req.records):
        _add_data_row(
            table,
            [
                r.get("date", ""),
                r.get("eventName", ""),
                r.get("staffName", ""),
                r.get("role", ""),
                r.get("scheduledStart", ""),
                r.get("scheduledEnd", ""),
                r.get("clockIn", ""),
                r.get("clockOut", ""),
                str(round(r.get("hoursWorked", 0), 1)),
                r.get("status", ""),
            ],
            idx,
            colors,
        )

    _add_summary_section(
        doc,
        [
            ("Total Records", str(req.summary.get("totalRecords", len(req.records)))),
            ("Total Hours", str(req.summary.get("totalHours", 0))),
        ],
        colors,
    )


_BUILDERS = {
    ReportType.STAFF_SHIFTS: _build_staff_shifts,
    ReportType.PAYROLL: _build_payroll,
    ReportType.ATTENDANCE: _build_attendance,
}


def _build_ai_analysis(doc: Document, req: ReportRequest, colors: dict[str, RGBColor] | None = None):
    """Render AI analysis markdown as Word paragraphs."""
    header_bg = (colors or {}).get("header_bg", _HEADER_BG)
    secondary = (colors or {}).get("secondary", RGBColor(0x33, 0x41, 0x55))
    md_text = ""
    if req.records and len(req.records) > 0:
        md_text = req.records[0].get("content", "")

    for line in md_text.split("\n"):
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph()
            continue
        if stripped.startswith("### "):
            h = doc.add_heading(stripped[4:], level=3)
            for run in h.runs:
                run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
        elif stripped.startswith("## "):
            h = doc.add_heading(stripped[3:], level=2)
            for run in h.runs:
                run.font.color.rgb = secondary
        elif stripped.startswith("# "):
            h = doc.add_heading(stripped[2:], level=1)
            for run in h.runs:
                run.font.color.rgb = header_bg
        elif stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(stripped[2:], style="List Bullet")
            for run in p.runs:
                run.font.size = Pt(10)
        else:
            # Handle **bold** segments
            p = doc.add_paragraph()
            parts = stripped.split("**")
            for i, part in enumerate(parts):
                if not part:
                    continue
                run = p.add_run(part)
                run.font.size = Pt(10)
                if i % 2 == 1:
                    run.font.bold = True
                    run.font.color.rgb = header_bg


def _build_working_hours(doc: Document, req: ReportRequest, colors: dict[str, RGBColor] | None = None):
    """Build working hours sheet table for Word."""
    header_bg = (colors or {}).get("header_bg", _HEADER_BG)
    # Event info block
    summary = req.summary or {}
    info_items = [
        ("Client", summary.get("client", "")),
        ("Event", summary.get("eventName", "")),
        ("Date", summary.get("date", "")),
        ("Schedule", f"{summary.get('startTime', '')} – {summary.get('endTime', '')}"),
        ("Venue", summary.get("venue", "")),
        ("Staff Count", str(len(req.records))),
    ]
    for label, value in info_items:
        if not value or value == " – ":
            continue
        p = doc.add_paragraph()
        run_label = p.add_run(f"{label}: ")
        run_label.font.bold = True
        run_label.font.size = Pt(10)
        run_label.font.color.rgb = header_bg
        run_val = p.add_run(str(value))
        run_val.font.size = Pt(10)

    doc.add_paragraph()

    cols = ["#", "Staff Name", "Role", "Phone", "ID", "Sched In", "Sched Out",
            "Clock In", "Clock Out", "Break", "Hours", "Signature"]
    table = doc.add_table(rows=1, cols=len(cols))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _style_header_row(table.rows[0], cols, colors)

    total_hours = 0.0
    for idx, r in enumerate(req.records):
        hours = r.get("totalHours", 0)
        total_hours += float(hours) if hours else 0
        _add_data_row(
            table,
            [
                str(idx + 1),
                r.get("name", ""),
                r.get("role", "Staff"),
                r.get("phone", ""),
                r.get("companyId", ""),
                r.get("scheduledIn", ""),
                r.get("scheduledOut", ""),
                r.get("clockIn", ""),
                r.get("clockOut", ""),
                r.get("breakDuration", ""),
                str(r.get("totalHours", "")),
                "",  # Signature — blank for manual entry
            ],
            idx,
            colors,
        )

    # Totals row
    totals_row = table.add_row()
    for i, cell in enumerate(totals_row.cells):
        if i == 0:
            cell.text = ""
        elif i == 1:
            cell.text = "TOTAL"
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.bold = True
                    run.font.size = Pt(9)
        elif i == 10:
            cell.text = str(round(total_hours, 1))
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.bold = True
                    run.font.size = Pt(9)
        else:
            cell.text = ""
        _set_cell_shading(cell, RGBColor(0xF1, 0xF5, 0xF9))

    # Sign-off section
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Manager Signature: ")
    run.font.size = Pt(10)
    run.font.bold = True
    p.add_run("_" * 40).font.size = Pt(10)

    p2 = doc.add_paragraph()
    run2 = p2.add_run("Client Representative: ")
    run2.font.size = Pt(10)
    run2.font.bold = True
    p2.add_run("_" * 40).font.size = Pt(10)


def _add_logo(doc: Document, logo_url: str) -> None:
    """Download logo from URL and insert into document header area."""
    try:
        with urlopen(logo_url) as resp:
            logo_data = resp.read()
        stream = BytesIO(logo_data)
        doc.add_picture(stream, width=Inches(2.0))
    except Exception:
        pass  # Logo fetch failed — continue without it


def create_report(req: ReportRequest, output_path: str) -> None:
    colors = _get_colors(req)
    theme = _get_theme_config(req)

    # Working hours uses its own builder
    if req.report_type == ReportType.WORKING_HOURS:
        builder = _build_working_hours
    # AI analysis uses a simplified prose builder
    elif req.report_type == ReportType.AI_ANALYSIS:
        builder = _build_ai_analysis
    else:
        builder = _BUILDERS.get(req.report_type)
        if not builder:
            raise ValueError(f"Unknown report type: {req.report_type}")

    doc = Document()

    # Logo (gated on theme)
    bc = req.brand_config
    if bc and bc.logo_header_url and theme.get("show_logo", True):
        _add_logo(doc, bc.logo_header_url)

    # Title — use primary color for classic/executive, dark for plain
    title_color = _hex_to_rgb((bc or BrandConfig()).primary_color) if req.template_design != TemplateDesign.PLAIN else RGBColor(0x1A, 0x1A, 0x1A)
    title_para = doc.add_heading(req.title, level=1)
    for run in title_para.runs:
        run.font.color.rgb = title_color

    # Subtitle with period
    subtitle = doc.add_paragraph()
    run = subtitle.add_run(f"{req.company_name}  |  {req.period.label}")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    doc.add_paragraph()

    builder(doc, req, colors)

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(
        f"Generated by {req.company_name} Document Service — "
        f"{datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}"
    )
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    doc.save(output_path)
